import re
import codecs
import chardet
import os
from docx import Document
from docx.oxml.ns import qn


class reformat_class():
    def switch_file(self, inputname):
        # 识别文件编码
        bytes = min(32, os.path.getsize(inputname))
        raw = open(inputname, 'rb').read(bytes)
        result = chardet.detect(raw)
        encoding = result['encoding']
        files = open(inputname, 'r', encoding=encoding)
        data = files.read()
        # utf-8 bom的文件转换
        if data != 'utf-8-sig':
            codecs.open(inputname, 'w', encoding='utf-8-sig').write(data)
        files.close()

    def reformat(self, inputname, outputname1):
        # 文件读取
        self.outputname1 = outputname1
        with open(inputname, encoding='utf-8-sig') as content:
            contents = content.read()
        # 数据处理，使用正则表达式
        exp = re.sub(r'(\])\n(.*)', r'\1 \2', contents)  # 使用分组方式替换字符串
        exp_sec = re.sub(r'\n\n', r'\n', exp)
        exp_thi = re.sub(r'(\)+)\n([a-zA-Z])', r'\1\n\t\t\2', exp_sec)
        exp_fou = re.sub(r'(\s)\n(.*)', r'\1 \2', exp_thi)
        exp_fif = re.sub(r'([\u4e00-\u9fa5]+)\n([a-zA-Z])', r'\1\n\t\t\2', exp_fou)
        exp_six = re.sub(r'([\u4e00-\u9fa5]+)\n([\u4e00-\u9fa5]+)', r'\1\n\t\t\2', exp_fif)
        global final
        final = exp_six

    def reformat_done(self):
        # 写入文件
        with open(self.outputname1, 'a', encoding='utf-8-sig') as content_reformat:
            content_reformat.write(final)

    def reformat_second(self, inputname, outputname):
        self.outputname = outputname
        with open(inputname, encoding='utf-8-sig') as content:
            contents = content.read()
        exp = re.sub(r'(\])\n(.*)', r'\1\n\t\2', contents)
        exp_sec = re.sub(r'\n\n', r'\n', exp)
        exp_thi = re.sub(r'(\s)\n(.*)', r'\n\t\2', exp_sec)
        exp_fou = re.sub(r'([\u4e00-\u9fa5]+)\n([a-zA-Z])', r'\1\n\t\2', exp_thi)
        exp_fif = re.sub(r'([\u4e00-\u9fa5]+)\n([\u4e00-\u9fa5]+)', r'\1\n\t\2', exp_fou)
        global final_second
        final_second = exp_fif

    def reofrmat_done_second(self):
        with open(self.outputname, 'a', encoding='utf-8-sig') as content_reformat:
            content_reformat.write(final_second)


class word_dispose(reformat_class):
    def word_create(self):
        self.document = Document()
        self.document.styles['Normal'].font.name = '等线'  # 设置字体
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    def word_write(self, word_name):
        self.document.add_paragraph(final)
        self.document.save(word_name)

    def word_write_second(self, word_name):
        self.document.add_paragraph(final_second)
        self.document.save(word_name)


final = ''
final_second = ''
# LINUX, MACOS 用户自行修改
inputname = 'word.txt'
outputname = 'word-reformated.txt'

if __name__ == "__main__":
    reformat = reformat_class()
    reformat.reformat_second(inputname, outputname)
